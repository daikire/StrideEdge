"""Phase 1 要件定義・設計思想 Word文書生成スクリプト（v1.1 レビュー反映版）"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

style_normal = doc.styles['Normal']
style_normal.font.name = 'Hiragino Kaku Gothic ProN'
style_normal.font.size = Pt(10.5)

def set_font(run, bold=False, size=None, color=None):
    run.font.name = 'Hiragino Kaku Gothic ProN'
    run.font.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.name = 'Hiragino Kaku Gothic ProN'
    p.runs[0].font.size = Pt(16)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def heading2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.name = 'Hiragino Kaku Gothic ProN'
    p.runs[0].font.size = Pt(13)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def heading3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.name = 'Hiragino Kaku Gothic ProN'
    p.runs[0].font.size = Pt(11)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def para(text, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.75)
    run = p.add_run(text)
    set_font(run, bold=bold)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.75 + level * 0.5)
    run = p.add_run(text)
    set_font(run)
    return p

def note(text):
    """注釈ボックス的な段落"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75)
    run = p.add_run('【注】' + text)
    set_font(run, size=9.5, color=(0x70, 0x70, 0x70))
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            set_font(run, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '2E74B5')
        tcPr.append(shd)
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row_data in enumerate(rows):
        row = table.rows[ri+1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = val
            for run in cell.paragraphs[0].runs:
                set_font(run)
            if ri % 2 == 1:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'DEEAF1')
                tcPr.append(shd)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

# =========================================================
# 表紙
# =========================================================
doc.add_paragraph()
doc.add_paragraph()
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('StrideEdge\n実データ連携 Phase 1\n要件定義書 / 設計思想書')
run.font.name = 'Hiragino Kaku Gothic ProN'
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

doc.add_paragraph()
meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta_p.add_run(
    f'バージョン: 1.1（レビュー反映版）\n'
    f'作成日: {datetime.date.today().strftime("%Y年%m月%d日")}\n'
    f'作成者: StrideEdge開発チーム\n'
    f'対象システム: StrideEdge 競馬予想支援アプリ'
)
set_font(run, size=11)
doc.add_page_break()

# =========================================================
# 1. ドキュメント概要
# =========================================================
heading1('1. ドキュメント概要')

heading2('1.1 目的')
para(
    '本書は StrideEdge における実データ連携 Phase 1 の要件定義および設計思想を定めるものである。'
    '現在StrideEdgeはダミーデータで動作しているが、本フェーズにより実際のレース情報・出走表・'
    'オッズをnetkeiba.comから自動取得し、スコアリング精度と実用性の向上を目指す。'
)
note('本実装は個人の学習環境のみを前提とする。外部公開・商用利用は禁止する。')

heading2('1.2 対象読者')
bullet('StrideEdge開発チームメンバー（実装担当・レビュー担当）')
bullet('外部AIレビュアー（設計妥当性チェック）')
bullet('将来の保守担当者')

heading2('1.3 スコープ')
para('本書が扱う範囲（Phase 1）:')
bullet('netkeiba.com からのレース情報・出走表スクレイピング実装（手動トリガーのみ）')
bullet('取得データのStrideEdge DBへの保存・マッピング')
bullet('フロントエンド /datasources 画面への取得状況表示・手動更新ボタン追加')
bullet('スクレイプログ管理（scrape_logs テーブル）')

para('本書が扱わない範囲（Phase 2 以降）:')
bullet('馬の過去成績の構造化取得・距離適性スコア改善（Phase 2）')
bullet('騎手成績動的算出・馬場適性スコア追加（Phase 3）')
bullet('オッズ再取得スケジューラー・結果自動取得（Phase 4）')
bullet('data_sourceカラム追加によるダミー/実データ区別UI（Phase 2以降）')
bullet('JRA-VAN 公式API連携（将来フェーズ）')

doc.add_page_break()

# =========================================================
# 2. 背景・課題
# =========================================================
heading1('2. 背景・課題')

heading2('2.1 現状の課題')
add_table(
    ['課題ID', '内容', '影響'],
    [
        ['KI-001', 'サンプルデータは架空データ。実際のレース情報との接続が未実装', '実運用不可'],
        ['─', 'オッズがダミー固定値のため、スコアリング精度が低い', 'スコア信頼性ゼロ'],
        ['─', '馬の過去成績が文字列のみ（距離・馬場情報なし）', '適性スコア計算不可（Phase 2で対応）'],
        ['─', 'レース結果は手動入力のみ', '運用コスト高（Phase 4で対応）'],
    ],
    col_widths=[2.5, 9.0, 4.5]
)

heading2('2.2 Phase 1で解決する課題')
bullet('実際のレース情報（開催日・会場・グレード・距離）を自動取得')
bullet('実際の出走表（馬名・騎手・枠順・斤量・単勝オッズ・人気）を自動取得')
bullet('取得状況をUIで可視化し、手動更新を可能にする')

doc.add_page_break()

# =========================================================
# 3. 機能要件
# =========================================================
heading1('3. 機能要件')

heading2('3.1 データ取得機能')

heading3('3.1.1 取得対象ページ')
para('netkeiba.comの以下のページを対象とする（個人利用・学習目的）:')
add_table(
    ['データ種別', '取得元ページ種別', '取得タイミング（Phase 1）'],
    [
        ['開催カレンダー', '開催一覧ページ（月別）', '手動トリガー時（1回取得）'],
        ['レース情報', 'レース一覧ページ（開催日別）', '手動トリガー時（1回取得）'],
        ['出走表', 'レース出走表ページ', '手動トリガー時（1回取得）'],
        ['単勝オッズ', 'オッズページ', '手動トリガー時（1回取得）'],
    ],
    col_widths=[3.5, 6.5, 5.5]
)
note(
    'Phase 1では自動スケジューリングは実装しない。'
    'スケジューラー（前日夜/当日朝の自動実行・発走30分前のオッズ再取得）はPhase 4のスコープ。'
    '推奨手動実行タイミング: 前日22時〜当日8時の間。'
)

heading3('3.1.2 race_idの採番規則（RV-P1-02反映）')
para('race_idはnetkeiba.comのレースID（URLパスに含まれる12桁数字）をそのまま使用する。')
add_table(
    ['項目', '内容', '例'],
    [
        ['採番方式', 'netkeiba レースID（12桁数字）をそのまま流用', '202504190101'],
        ['フォーマット', 'YYYYMMDDVVRR（年月日・会場コード・レース番号）', '20250419: 2025年4月19日'],
        ['horse_idとの統一', 'horse_idも同様にnetkeiba IDを流用（既存設計と統一）', '─'],
        ['将来の移行', '内部IDへの移行が必要になった場合はマイグレーションで対応', '─'],
    ],
    col_widths=[3.5, 8.0, 4.0]
)

heading3('3.1.3 取得データ項目マッピング')
add_table(
    ['DBテーブル', 'カラム', 'NULL可否', 'netkeiba取得元', '備考'],
    [
        ['races', 'race_id', 'NOT NULL', 'URL内の12桁レースID', '必須（UPSERTキー）'],
        ['races', 'race_name', 'NOT NULL', 'レース名', '必須フィールド'],
        ['races', 'race_date', 'NOT NULL', '開催日', 'YYYY-MM-DD形式。必須'],
        ['races', 'venue', 'NOT NULL', '開催場所', '中山/東京/阪神 等。必須'],
        ['races', 'race_number', 'NOT NULL', 'レース番号', '1R〜12R。必須'],
        ['races', 'distance', 'NOT NULL', 'レース距離', '数値（m）。必須'],
        ['races', 'surface', 'NOT NULL', '馬場種別', 'turf/dirt。必須'],
        ['races', 'grade', 'NULL可', 'グレード', 'G1/G2/G3等。未取得時NULL'],
        ['races', 'race_class', 'NULL可', 'クラス', '3歳未勝利等。未取得時NULL'],
        ['races', 'prize_money', 'NULL可', '本賞金', '取得可能な場合のみ'],
        ['horses', 'horse_id', 'NOT NULL', 'URL内のnetkeiba馬ID', 'UPSERTキー'],
        ['horses', 'horse_name', 'NOT NULL', '馬名', '必須'],
        ['horses', 'age', 'NULL可', '馬齢', '数値'],
        ['horses', 'sex', 'NULL可', '性別', '牡/牝/セン'],
        ['horses', 'trainer', 'NULL可', '調教師名', ''],
        ['horses', 'owner', 'NULL可', '馬主名', '取得できた場合のみ'],
        ['entries', 'horse_number', 'NOT NULL', '馬番', '必須'],
        ['entries', 'gate_number', 'NOT NULL', '枠番', '必須'],
        ['entries', 'jockey', 'NOT NULL', '騎手名', '必須'],
        ['entries', 'weight_carried', 'NULL可', '斤量', '数値（kg）'],
        ['entries', 'odds', 'NULL可', '単勝オッズ', '取得時点の値。発走前は変動する'],
        ['entries', 'popularity', 'NULL可', '人気順位', '数値'],
        ['entries', 'horse_weight', 'NULL可', '馬体重', '数値（kg）。前日確定後のみ取得可'],
        ['entries', 'horse_weight_diff', 'NULL可', '馬体重増減', '数値（kg）±'],
        ['entries', 'recent_results', 'NULL可', '直近着順', '文字列形式（下記参照）'],
    ],
    col_widths=[2.5, 3.5, 2.0, 4.0, 3.5]
)

heading3('3.1.4 recent_resultsのフォーマット定義（RV-P1-04反映）')
para('Phase 1のrecent_resultsは「直近5走の着順文字列（表示用途のみ）」とし、以下のフォーマットで保存する:')
bullet('フォーマット: 着順をハイフン区切りで結合（例: "1-3-2-5-1"）')
bullet('最大5走分。取得できた件数分のみ記録（例: 過去2走のみなら "1-3"）')
bullet('中止・除外の場合は "中" または "除" として記録')
bullet('構造化データ（距離・馬場別の実績）としての過去成績はPhase 2のスコープ')
note('Phase 2でhorse_race_historyテーブルを追加し、距離・馬場別の詳細実績を管理する予定。')

heading3('3.1.5 NULL保存ポリシー（RV-P1-03反映）')
para('パース失敗時のフィールドに対するNULL許容ポリシーを以下のとおり定める:')
add_table(
    ['カテゴリ', '対象フィールド', 'NULLの場合の処理'],
    [
        ['必須フィールド', 'race_id, race_name, race_date, venue, race_number, distance, surface\nhorse_id, horse_name\nhorse_number, gate_number, jockey',
         'そのレース/馬/エントリーレコードを保存せず、scrape_logsにpartialを記録する'],
        ['任意フィールド', 'grade, race_class, prize_money, age, sex, trainer, owner\nweight_carried, odds, popularity, horse_weight, horse_weight_diff, recent_results',
         'NULLのまま保存する。scrape_logsにはpartialではなくsuccessを記録'],
    ],
    col_widths=[3.0, 7.0, 5.5]
)

heading2('3.2 スクレイプログ管理機能')
para('取得の成否・日時・対象URLを記録するscrape_logsテーブルを追加する:')
add_table(
    ['カラム名', '型', 'NULL可否', '内容'],
    [
        ['id', 'INTEGER PK', 'NOT NULL', '自動採番'],
        ['target_date', 'TEXT', 'NOT NULL', '対象開催日（YYYY-MM-DD）'],
        ['url', 'TEXT', 'NOT NULL', '取得対象URL'],
        ['status', 'TEXT', 'NOT NULL', 'success / partial / error / running / block_detected'],
        ['races_fetched', 'INTEGER', 'NULL', '取得成功レース数'],
        ['entries_fetched', 'INTEGER', 'NULL', '取得成功エントリー数'],
        ['error_message', 'TEXT(1000)', 'NULL', 'エラー概要（1000文字以内）。詳細はscraper.logに出力'],
        ['scraped_at', 'TIMESTAMP', 'NOT NULL', '取得実行日時（CURRENT_TIMESTAMP）'],
    ],
    col_widths=[3.5, 2.8, 2.0, 7.2]
)
note(
    'statusにrunningを追加し、同時実行中のsyncジョブが存在する場合は'
    '後続リクエストを409 Conflictで弾く「二重実行防止」に使用する。'
    'block_detectedはIPブロック検知時に記録する（後述）。'
)

heading2('3.3 フロントエンド拡張機能（/datasources 画面）')
para('/datasources 画面に以下を追加する（Phase 1スコープ）:')
bullet('直近スクレイプ履歴の一覧表示（日付・取得件数・ステータス・実行日時）')
bullet('手動でスクレイピングを実行するボタン（対象日を指定可能）')
bullet('取得失敗時のエラーメッセージ表示')
para('以下はPhase 1スコープ外（Phase 2で対応）:')
bullet('実データ/ダミーデータのタグ表示（data_sourceカラムの追加がPhase 2のため）')

doc.add_page_break()

# =========================================================
# 4. 非機能要件
# =========================================================
heading1('4. 非機能要件')

heading2('4.1 パフォーマンス')
add_table(
    ['項目', '要件', '補足'],
    [
        ['1レースあたりの取得時間', '10秒以内', 'タイムアウト設定10s。AC-09で検証'],
        ['1開催日（全レース）の取得時間', '3分以内', '12レース想定。AC-09で検証'],
        ['リクエスト間隔', '最低2秒以上', 'サーバー負荷軽減・礼儀的クローリング'],
        ['1セッションの最大リクエスト数', '150リクエスト以内', 'IPブロック対策。超過時は処理中断'],
        ['同時リクエスト数', '1（逐次処理）', '並列スクレイピング禁止'],
    ],
    col_widths=[5.5, 3.5, 6.5]
)

heading2('4.2 信頼性・エラーハンドリング')
bullet('HTTP 429 / 5xx 発生時: 指数バックオフでリトライ（最大3回。1回目5秒・2回目15秒・3回目45秒）')
bullet('HTTP 403 / 予期しないリダイレクト発生時: IPブロックと判断し処理を即座に中断。scrape_logsにblock_detectedを記録')
bullet('必須フィールドのパース失敗時: そのレコードを保存せず、scrape_logsにpartialを記録')
bullet('任意フィールドのパース失敗時: NULLで保存し、scrape_logsはsuccessのまま')
bullet('全レース取得失敗時: scrape_logsにerrorを記録し、既存ダミーデータで動作継続')
bullet('HTMLの構造変更検知: 主要CSSセレクタが取得できない場合にERRORログを出力')
bullet('同時実行防止: scrape_logsにrunningレコードが存在する場合は409 Conflictを返す')

heading2('4.3 セキュリティ・利用方針')
bullet('個人利用・学習目的に限定する')
bullet('netkeiba.comの利用規約およびrobots.txtを事前確認し、禁止パスへのアクセスは行わない')
bullet('User-Agentヘッダを適切に設定する（Pythonデフォルトのまま。ブラウザ偽装は行わない）')
bullet('取得データの商用利用・再配布は行わない')
bullet('本実装コードをpublic repositoryに公開する場合は利用規約の再確認が必要')
note(
    'netkeibaの利用規約にスクレイピング禁止条項が含まれる可能性がある。'
    '実装前に利用規約を確認し、禁止されている場合はJRA公式データ・JRA-VANへの移行を優先する。'
    '法的リスクが高いと判断された場合、Phase 1のデータソースをJRA公式サイトに変更する。'
)

heading2('4.4 SQLiteの設定（RV-P1-10反映）')
bullet('WALモード（PRAGMA journal_mode=WAL）を起動時に有効化する')
bullet('これにより読み取りと書き込みが競合せず、連打・将来のスケジューラー追加時にも安定動作する')
bullet('init_db.py の初期化処理に PRAGMA設定を追加する')

heading2('4.5 ログ・監視')
bullet('スクレイピング実行ログを logs/scraper.log に出力')
bullet('INFO: 取得開始・完了・件数・処理時間')
bullet('WARNING: パース失敗・部分取得・リトライ発生')
bullet('ERROR: タイムアウト・HTTP エラー・リトライ上限超過・IPブロック検知')

doc.add_page_break()

# =========================================================
# 5. 設計思想
# =========================================================
heading1('5. 設計思想')

heading2('5.1 基本方針')
para(
    'Phase 1の設計は「既存コードへの影響を最小化しながら実データを差し込める構造」を最優先とする。'
    'スクレイパーは独立したモジュールとして分離し、サービス層・DB層は既存インターフェースを維持したまま'
    '拡張する。将来のデータソース差し替え（JRA-VAN移行等）を考慮し、スクレイパーと保存ロジックを'
    '明確に分離する。'
)

heading2('5.2 モジュール分割方針（RV-P1-05・06反映）')
add_table(
    ['モジュール', 'パス', '責務'],
    [
        ['NetkeibaScraper', 'backend/app/scrapers/netkeiba_scraper.py',
         'HTML取得（fetch）とパース（parse）の分離。生データ（ScrapedRace/ScrapedEntry）を返す。DB構造を知らない'],
        ['DataSyncService', 'backend/app/services/data_sync_service.py',
         '処理のオーケストレーションのみ担当。NetkeibaScraperを呼び出し、RepositoryにDB保存を委譲'],
        ['RaceRepository', 'backend/app/repositories/race_repository.py',
         'races / horses / entries テーブルへのUPSERT処理。scrape_logs への記録'],
        ['SyncRouter', 'backend/app/routers/sync.py',
         'POST /api/sync/races・GET /api/sync/logs エンドポイント定義'],
        ['init_db（拡張）', 'backend/app/database/init_db.py',
         'scrape_logsテーブル追加・WALモード設定追加'],
    ],
    col_widths=[3.5, 6.0, 6.0]
)
note(
    'DataSyncServiceはスクレイパーの具体実装（netkeiba等）を直接知らない構造にする。'
    '将来JRA-VANへ差し替える際はNetkeibaScraperをJraVanClientに置き換えるだけでよい。'
)

heading2('5.3 スクレイパー設計原則')

heading3('5.3.1 HTMLの取得とパースの分離')
para(
    'NetkeibaScraperは「HTML文字列の取得（fetch_html）」と「HTMLのパース（parse_race_list等）」を'
    '別メソッドに分離する。これによりユニットテスト時はHTMLファイルをモックとして使用でき、'
    'ネットワーク依存のテストを避けられる。'
)

heading3('5.3.2 生データとDBモデルの分離')
para(
    'スクレイパーはPydanticモデル（ScrapedRace, ScrapedEntry等）を返す。'
    'DBへのマッピングはRaceRepositoryが担う。スクレイパーはDB構造を知らない。'
    'これによりHTMLの変更時にスクレイパーのみ修正すればよい。'
)

heading3('5.3.3 冪等性の確保（RV-P1-07反映）')
para(
    '同じ日のデータを複数回取得しても重複登録されないよう、races/horses/entriesへの保存は'
    'INSERT ... ON CONFLICT(id) DO UPDATE SET（SQLite 3.24以降）を使用する。'
)
para(
    'INSERT OR REPLACEは使用しない。理由: INSERT OR REPLACEはPK衝突時に既存行をDELETE→INSERT'
    'するため、NULLが混入したデータで既存の正常データが上書きされるリスクがある。'
    'ON CONFLICT DO UPDATE SETは更新対象カラムを明示的に指定できるため安全。'
)

heading3('5.3.4 IPブロック検知（RV-P1-08反映）')
para('以下の条件を検知した場合にIPブロックと判断し、処理を中断する:')
bullet('HTTPステータスコード 403')
bullet('リダイレクト先URLがログインページや認証ページになっている場合')
bullet('レスポンスボディに「アクセスが制限」等のキーワードが含まれる場合')
para('IPブロック検知時の処理:')
bullet('scrape_logsにblock_detectedステータスを記録')
bullet('ERROR ログを logs/scraper.log に出力')
bullet('APIレスポンスに block_detected フラグを含めてフロントエンドに通知')

heading2('5.4 データフロー')
doc.add_paragraph()
flow_p = doc.add_paragraph()
flow_p.paragraph_format.left_indent = Cm(1.0)
run = flow_p.add_run(
    '① フロントエンド\n'
    '    POST /api/sync/races?date=YYYY-MM-DD\n'
    '        ↓\n'
    '② SyncRouter\n'
    '    実行中チェック（scrape_logsにrunning存在→409 Conflict）\n'
    '    DataSyncService.sync_races(date) を async で呼び出し\n'
    '        ↓\n'
    '③ DataSyncService（オーケストレーション）\n'
    '    scrape_logsに status=running で記録\n'
    '    await NetkeibaScraper.fetch_race_list(date) → ScrapedRace[]\n'
    '    各レースに対して await NetkeibaScraper.fetch_entries(race_url) → ScrapedEntry[]\n'
    '        ↓\n'
    '④ RaceRepository（DB保存）\n'
    '    races / horses / entries テーブルに ON CONFLICT DO UPDATE\n'
    '        ↓\n'
    '⑤ DataSyncService\n'
    '    scrape_logsを status=success/partial/error に更新\n'
    '        ↓\n'
    '⑥ SyncRouter\n'
    '    { status, races_fetched, entries_fetched, errors[] } を返却\n'
    '        ↓\n'
    '⑦ フロントエンド\n'
    '    /datasources 画面に取得結果を表示'
)
set_font(run, size=10)

heading2('5.5 既存機能への影響')
para(
    '既存の分析・買い目提案ロジックはDBのracesおよびentriesテーブルを参照している。'
    'Phase 1ではこれらのテーブルに実データを書き込むだけでよく、analysis_service.pyや'
    'ticket_service.pyへの変更は不要。スコアリングの精度は実データ投入により自動的に向上する。'
)

doc.add_page_break()

# =========================================================
# 6. 新規API仕様
# =========================================================
heading1('6. 新規APIエンドポイント仕様')

heading2('6.1 POST /api/sync/races')
add_table(
    ['項目', '内容'],
    [
        ['メソッド', 'POST'],
        ['パス', '/api/sync/races'],
        ['クエリパラメータ', 'date: string（YYYY-MM-DD）必須'],
        ['409 Conflict', '{ "detail": "既にsync実行中です" }（同時実行防止）'],
        ['200 success', '{ "status": "success", "races_fetched": 12, "entries_fetched": 144, "errors": [] }'],
        ['200 partial', '{ "status": "partial", "races_fetched": 8, "entries_fetched": 96, "errors": ["R04: 必須フィールド取得失敗"] }'],
        ['200 block_detected', '{ "status": "block_detected", "message": "IPブロックを検知しました。時間をおいて再試行してください" }'],
        ['500 error', '{ "status": "error", "message": "ネットワークエラー: ..." }'],
    ],
    col_widths=[4.0, 11.5]
)

heading2('6.2 GET /api/sync/logs')
add_table(
    ['項目', '内容'],
    [
        ['メソッド', 'GET'],
        ['パス', '/api/sync/logs'],
        ['クエリパラメータ', 'limit: int（デフォルト20）'],
        ['レスポンス', 'ScrapeLog[] の配列（scraped_at降順）'],
    ],
    col_widths=[4.0, 11.5]
)

doc.add_page_break()

# =========================================================
# 7. DB拡張仕様
# =========================================================
heading1('7. DB拡張仕様')

heading2('7.1 新規テーブル: scrape_logs')
para('既存テーブルへの構造変更はなし。scrape_logsテーブルのみを追加する。')
add_table(
    ['カラム', '型', 'NULL可否', 'デフォルト', '説明'],
    [
        ['id', 'INTEGER', 'NOT NULL', 'AUTOINCREMENT', '主キー'],
        ['target_date', 'TEXT', 'NOT NULL', '─', '対象開催日（YYYY-MM-DD）'],
        ['url', 'TEXT', 'NOT NULL', '─', '取得対象URL'],
        ['status', 'TEXT', 'NOT NULL', '─', 'success / partial / error / running / block_detected'],
        ['races_fetched', 'INTEGER', 'NULL', '0', '取得レース数'],
        ['entries_fetched', 'INTEGER', 'NULL', '0', '取得エントリー数'],
        ['error_message', 'TEXT(1000)', 'NULL', 'NULL', 'エラー概要（1000文字以内）。詳細はscraper.log'],
        ['scraped_at', 'TIMESTAMP', 'NOT NULL', 'CURRENT_TIMESTAMP', '実行日時'],
    ],
    col_widths=[3.5, 2.8, 2.0, 2.5, 4.7]
)

heading2('7.2 既存テーブルの変更')
para('Phase 1では既存テーブル（races/horses/entries）への構造変更は行わない。')
para('data_source カラム（"dummy" / "real"）の追加はPhase 2のスコープとする。')

heading2('7.3 SQLite設定変更')
bullet('init_db.py の接続時に PRAGMA journal_mode=WAL を実行する')
bullet('対象: アプリ起動時の初期化処理（既存のinit_db関数に追加）')

doc.add_page_break()

# =========================================================
# 8. 受け入れ条件
# =========================================================
heading1('8. 受け入れ条件（Definition of Done）')

heading2('8.1 機能面（RV-P1-11反映: 検証手順・合格基準を明記）')
add_table(
    ['No.', '受け入れ条件', '検証手順', '合格基準'],
    [
        ['AC-01',
         'POST /api/sync/races?date=YYYY-MM-DDを実行すると、当日のレース情報がDBに保存される',
         'sync後に SELECT COUNT(*) FROM races WHERE race_date=対象日 を実行',
         'netkeibaの当該日レース件数と一致すること'],
        ['AC-02',
         '出走馬・オッズ・騎手・枠順がDBに保存される',
         'SELECT COUNT(*) FROM entries WHERE race_id IN (対象レースID群)',
         '出走表の全エントリー数と一致すること'],
        ['AC-03',
         '同じ日を2回syncしても重複レコードが作成されない',
         '同一dateで2回POST実行。前後のSELECT COUNT(*)を比較',
         '2回目実行後もレコード件数が変化しないこと'],
        ['AC-04',
         'scrape_logsに実行日時・ステータス・取得件数が記録される',
         'sync後にSELECT * FROM scrape_logs ORDER BY scraped_at DESC LIMIT 1',
         'status/races_fetched/entries_fetched/scraped_atが正しく記録されていること'],
        ['AC-05',
         '/datasources画面にscrape_logsの一覧が表示される',
         'ブラウザで/datasourcesを開く',
         'scrape_logsの内容（日付・件数・ステータス）が一覧表示されること'],
        ['AC-06',
         '手動更新ボタンを押すと取得が実行される',
         'ブラウザで対象日を指定してボタンを押下',
         'scrape_logsに新しいレコードが追加され、画面に結果が表示されること'],
        ['AC-07',
         'ネットワークエラー時にアプリがクラッシュしない',
         'ネットワーク遮断状態でPOST /api/sync/racesを実行',
         '500エラーが返り、ログにERRORが記録され、アプリは正常動作を継続すること'],
        ['AC-08',
         '/races画面にsyncした実データが表示される',
         'sync後にブラウザで/racesを開き、対象日を選択',
         'syncで取得したレース名・馬名・オッズが表示されること'],
        ['AC-09',
         '性能要件を満たす',
         'sync実行時の開始〜完了時間をAPIレスポンスのログで計測',
         '1レース分: 10秒以内。全レース（最大12レース）: 3分以内'],
        ['AC-10',
         '同時実行防止が機能する',
         'sync実行中に2回目のPOST /api/sync/racesを送信',
         '409 Conflictが返ること'],
    ],
    col_widths=[1.0, 4.0, 5.5, 5.0]
)

heading2('8.2 コード品質面')
bullet('NetkeibaScraperのparse系メソッドに対するユニットテストが存在する（HTMLファイルをモック使用）')
bullet('pytest でエラーなく通過する')
bullet('既存のエンドポイントが正常に動作することを確認（リグレッションなし）')

doc.add_page_break()

# =========================================================
# 9. 未解決事項・リスク
# =========================================================
heading1('9. 未解決事項・リスク（RV-P1-08・09・10反映）')

add_table(
    ['No.', 'リスク', 'レベル', '対応方針'],
    [
        ['R-01',
         'netkeibaのHTMLが変更された場合、パースが壊れる',
         '中',
         'パース失敗をERRORログに記録し即時検知。CSSセレクタを定数（SELECTORS辞書）で管理し修正を局所化'],
        ['R-02',
         '利用規約・robots.txtの変更によりアクセス制限が生じる',
         '中',
         '実装前にnetkeibaの利用規約を確認。禁止されている場合はJRA公式サイト/JRA-VANへ移行'],
        ['R-03',
         'IPブロックによるアクセス制限',
         '高',
         'リクエスト間隔2秒以上・1セッション150リクエスト上限・IPブロック検知ロジックを実装。検知時は即座に中断'],
        ['R-04',
         'オッズはリアルタイムで変動するため取得タイミングにより精度が変わる',
         '低',
         'Phase 1は手動1回取得のみ。スケジューラーによる再取得はPhase 4で対応'],
        ['R-05',
         '馬IDのnetkeiba依存',
         '低',
         'horse_idはnetkeiba IDを流用（race_idと同方針）。将来の移行はDBマイグレーションで対応可能な設計とする'],
        ['R-06',
         'SQLiteの並行書き込みによるdatabase is lockedエラー',
         '低（現状手動のみ）',
         'WALモード有効化（PRAGMA journal_mode=WAL）および同時実行防止（409 Conflict）で対処'],
    ],
    col_widths=[0.8, 5.0, 2.0, 7.7]
)

doc.add_page_break()

# =========================================================
# 10. 変更履歴
# =========================================================
heading1('10. 変更履歴')

add_table(
    ['バージョン', '日付', '変更内容', '担当'],
    [
        ['1.0', datetime.date.today().strftime('%Y/%m/%d'), '初稿作成', 'StrideEdge開発チーム'],
        ['1.1', datetime.date.today().strftime('%Y/%m/%d'),
         'レビュー（RV-P1-01〜15）反映: race_id採番規則追加・NULL保存ポリシー明確化・'
         'UPSERTをON CONFLICT DO UPDATEに変更・IPブロック対策追加・'
         '法的リスク追記・WALモード追加・AC検証基準追加・AC-09/10追加・'
         'オッズ取得タイミングをPhase 1は手動1回のみに修正・'
         'data_sourceタグ表示をPhase 2に移動・クラス名をNetkeibaScraperに修正',
         'StrideEdge開発チーム'],
    ],
    col_widths=[2.0, 3.0, 10.0, 3.5]
)

output_path = '/Users/Daiki/ClaudeCode/030_StrideEdge/01_pm/Phase1_要件定義書_設計思想書.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
