#!/usr/bin/env python3
"""GEMINI.md → GEMINI.docx 変換スクリプト"""

import sys
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def convert(src: str, dst: str) -> None:
    doc = Document()

    # デフォルトフォントを設定
    style = doc.styles["Normal"]
    style.font.name = "Hiragino Sans"
    style.font.size = Pt(10.5)

    with open(src, encoding="utf-8") as f:
        lines = f.readlines()

    in_code_block = False
    code_lines = []

    for raw in lines:
        line = raw.rstrip("\n")

        # コードブロック開始/終了
        if line.startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                in_code_block = False
                para = doc.add_paragraph()
                run = para.add_run("\n".join(code_lines))
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                para.paragraph_format.left_indent = Pt(20)
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        # 見出し
        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            text = m.group(2)
            doc.add_heading(text, level=level)
            continue

        # 水平線
        if re.match(r"^-{3,}$", line.strip()):
            doc.add_paragraph("─" * 40)
            continue

        # テーブル行（簡易対応）
        if line.startswith("|") and line.endswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            # 区切り行はスキップ
            if all(re.match(r"^[-: ]+$", c) for c in cells if c):
                continue
            table = doc.add_table(rows=1, cols=len(cells))
            table.style = "Table Grid"
            row = table.rows[0]
            for i, cell_text in enumerate(cells):
                # マークダウン装飾を除去
                cell_text = re.sub(r"\*{1,2}(.*?)\*{1,2}", r"\1", cell_text)
                row.cells[i].text = cell_text
            continue

        # 空行
        if not line.strip():
            doc.add_paragraph()
            continue

        # 箇条書き
        if re.match(r"^[-*]\s+", line):
            text = re.sub(r"^[-*]\s+", "", line)
            text = re.sub(r"\*{1,2}(.*?)\*{1,2}", r"\1", text)
            text = re.sub(r"`(.*?)`", r"\1", text)
            doc.add_paragraph(text, style="List Bullet")
            continue

        # 番号付きリスト
        if re.match(r"^\d+\.\s+", line):
            text = re.sub(r"^\d+\.\s+", "", line)
            text = re.sub(r"\*{1,2}(.*?)\*{1,2}", r"\1", text)
            doc.add_paragraph(text, style="List Number")
            continue

        # 通常テキスト（インライン装飾を除去）
        text = re.sub(r"\*{1,2}(.*?)\*{1,2}", r"\1", line)
        text = re.sub(r"`(.*?)`", r"\1", text)
        doc.add_paragraph(text)

    doc.save(dst)
    print(f"Converted: {src} → {dst}")

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: md_to_docx.py <input.md> <output.docx>")
        sys.exit(1)
    convert(sys.argv[1], sys.argv[2])
